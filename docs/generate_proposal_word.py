"""
Generator Dokumen Word — Proposal Lengkap Genesis (Subbab 1.1 s.d. 1.6)
Sesuai Standar LKS Nasional 2026, KBBI Baku, Margin 4-3-3-3, 1.5 Spasi, TNR 12, Tepat 4 Halaman A4
Fokus Utama: Aksi Iklim & Kerusakan Lingkungan (Tanpa Kata 'Sampah' guna Menghindari Blunder Akademik)
Teknologi: Google GenAI SDK (Vertex AI), Vertex AI Search RAG, Native Speech-to-Text (Tanpa OpenRouter / Whisper-1)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = r"d:\PROJECT ARIEF\LKS Dikdasmen\docs\Proposal_Teknis_Genesis.docx"

# Konstanta Format Dokumen
FONT_NAME = "Times New Roman"
SIZE_BODY = Pt(12)
SIZE_HEADING = Pt(12)
SIZE_TABLE_TEXT = Pt(10)
LINE_SPACING_BODY = 1.5
LINE_SPACING_TABLE = 1.15

COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY = RGBColor(0x33, 0x33, 0x33)

HEX_BORDER = "CCCCCC" # Warna batas abu-abu terang untuk tabel sederhana
HEX_HEADER_BG = "F2F2F2" # Latar belakang header abu-abu terang

def style_cell(cell, top_padding=60, bottom_padding=60, left_padding=100, right_padding=100, 
               shading_color=None, border_color=HEX_BORDER, v_align="center"):
    """Mengatur padding internal rapat (3pt atas/bawah), shading, border grid tipis, dan perataan vertikal tengah."""
    tcPr = cell._tc.get_or_add_tcPr()
    
    # 1. Padding Sel Rapat (60 dxa = 3pt atas/bawah)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top_padding}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom_padding}" w:type="dxa"/>'
        f'  <w:left w:w="{left_padding}" w:type="dxa"/>'
        f'  <w:right w:w="{right_padding}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)
    
    # 2. Shading Sel
    if shading_color:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
        tcPr.append(shd)
        
    # 3. Border Grid Sederhana (Semua sisi)
    for el in list(tcPr):
        if el.tag.endswith('tcBorders'):
            tcPr.remove(el)
    borders_xml = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders_xml)
    
    # 4. Perataan Vertikal Tengah (Vertical Centering)
    for el in list(tcPr):
        if el.tag.endswith('vAlign'):
            tcPr.remove(el)
    v_align_xml = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{v_align}"/>')
    tcPr.append(v_align_xml)

def write_cell(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING_TABLE
    
    # Memproses format miring otomatis pada istilah bahasa Inggris di dalam sel
    words = text.split(" ")
    for word in words:
        run = p.add_run(word + " ")
        run.font.name = FONT_NAME
        run.font.size = SIZE_TABLE_TEXT
        run.font.color.rgb = COLOR_BLACK
        run.bold = bold
        
        # Deteksi istilah asing/teknis untuk otomatis dicetak miring
        is_english = italic or any(term in word.lower() for term in [
            "multipart", "in-memory", "blur", "embedding", "streaming", "routing", 
            "check_duplicate_report", "pgvector", "hnsw", "whisper", "client", "record", 
            "geolocator", "database", "vision", "sharp", "flash", "openrouter", "throttler", 
            "jwt", "rbac", "rate", "limiting", "injection", "evasion", "typoglycemia", 
            "active", "learning", "feedback", "loop", "upvote", "design", "hitl", "gateway",
            "support", "similarity", "pending_human", "vertex", "search", "genai", "sdk"
        ])
        run.italic = is_english

        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
            rPr.append(rFonts)

def add_heading(doc, num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6) # Spasi sebelum heading diperketat
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    pf.keep_with_next = True
    
    run = p.add_run(f"{num} {title}")
    run.font.name = FONT_NAME
    run.font.size = SIZE_HEADING
    run.font.color.rgb = COLOR_BLACK
    run.bold = True

def add_paragraph(doc, segments, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after # Spasi antar paragraf diperketat
    pf.line_spacing = LINE_SPACING_BODY
    pf.first_line_indent = Cm(1.27)
    
    for txt, b, i in segments:
        run = p.add_run(txt)
        run.font.name = FONT_NAME
        run.font.size = SIZE_BODY
        run.font.color.rgb = COLOR_BLACK
        run.bold = b
        run.italic = i
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0
    
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True

def build_simple_table(doc, headers, data, col_widths):
    """Membangun tabel sederhana dengan kolom NO dan pembatas lengkap."""
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Header Row
    for col_idx, text in enumerate(headers):
        cell = tbl.rows[0].cells[col_idx]
        write_cell(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_cell(cell, shading_color=HEX_HEADER_BG)
    
    # Data Rows
    for row_idx, row_data in enumerate(data):
        row = tbl.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            
            # Kolom NO rata tengah, kolom lainnya rata kiri
            bold = (col_idx == 0) or (col_idx == 1)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            
            write_cell(cell, val, bold=bold, align=align)
            style_cell(cell)
                
    # Atur Lebar Kolom (Total tepat 14.0 cm untuk menjaga kelurusan dengan margin)
    for row in tbl.rows:
        for col_idx, width in enumerate(col_widths):
            row.cells[col_idx].width = Cm(width)
            
    return tbl


# ══════════════════════════════════════════════════════════════
# DOKUMEN SETUP (Margin 4-3-3-3 Sesuai Standar LKS)
# ══════════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(3.0)      # Atas: 3 cm
sec.bottom_margin = Cm(3.0)   # Bawah: 3 cm
sec.left_margin = Cm(4.0)     # Kiri: 4 cm (untuk jilid)
sec.right_margin = Cm(3.0)    # Kanan: 3 cm

# ══════════════════════════════════════════════════════════════
# PAGE 1: HEADER IDENTITAS & 1.1 LATAR BELAKANG
# ══════════════════════════════════════════════════════════════
# Header Identitas Tim di bagian paling atas Halaman 1
header_table = doc.add_table(rows=1, cols=1)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
header_table.autofit = False
header_table.rows[0].cells[0].width = Cm(14.0)

cell = header_table.rows[0].cells[0]
style_cell(cell, top_padding=80, bottom_padding=80, left_padding=120, right_padding=120, shading_color="F2F2F2")

# Tulis identitas terstruktur di dalam kotak abu-abu premium
write_cell(cell, "PROPOSAL SOLUSI (PROBLEM CANVAS) — EKSHIBISI AI LKS NASIONAL 2026\n"
                 "Nama Platform: GENESIS | Nama Tim: [Nama Tim Anda]\n"
                 "Asal Sekolah: [Nama Sekolah Anda]\n"
                 "Anggota: 1. [Anggota 1] 2. [Anggota 2] 3. [Anggota 3] 4. [Anggota 4] 5. [Anggota 5]\n"
                 "Guru Pembimbing: [Nama Guru Pendamping]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# Jarak minimal setelah header
spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(4)
spacer.paragraph_format.space_after = Pt(0)
spacer.paragraph_format.line_spacing = 1.0

add_heading(doc, "1.1", "Latar Belakang dan Analisis Permasalahan")

add_paragraph(doc, [
    ("Kerusakan akibat perubahan iklim di Indonesia semakin kompleks dan saling berkaitan, memengaruhi ketahanan pangan, sumber daya air, "
     "kesehatan, hingga ekonomi. Kondisi ini menuntut pengambilan keputusan yang mampu menghubungkan informasi lintas sektor "
     "untuk kebijakan yang lebih tepat. Seiring transformasi digital, partisipasi masyarakat dalam pelaporan masalah dan kontribusi data "
     "terus meningkat. Didukung perkembangan kecerdasan artifisial dan infrastruktur data, informasi publik ini berkembang menjadi "
     "aset strategis dalam mendukung keputusan berbasis bukti.", False, False)
])

add_paragraph(doc, [
    ("Meskipun kemampuan mengolah data berkembang, sebagian besar sistem masih berfokus pada penyediaan informasi sebagai hasil akhir, "
     "bukan sebagai awal pembentukan pengetahuan. Laporan masyarakat masih tersimpan sebagai entitas terisolasi, sehingga hubungan antarperistiwa, "
     "pola berulang, dan pembelajaran akumulatif belum terbentuk optimal. Akibatnya, setiap permasalahan baru sering diperlakukan sebagai "
     "kasus terpisah, sementara pengalaman dan pengetahuan masa lalu belum dapat dijadikan landasan bagi keputusan berikutnya.", False, False)
])

add_paragraph(doc, [
    ("Tantangan utama pengelolaan informasi kini adalah membangun proses pembentukan pengetahuan secara berkelanjutan. Tanpa mekanisme "
     "yang menghubungkan dan memperkaya informasi untuk menghasilkan pembelajaran, penambahan volume data hanya memperbesar akumulasi "
     "informasi tanpa meningkatkan kualitas pemahaman. Oleh karena itu, diperlukan pendekatan yang mampu menghubungkan informasi secara "
     "berkelanjutan sebagai fondasi bagi sistem pengambilan keputusan yang adaptif, kolaboratif, dan berbasis bukti.", False, False)
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PAGE 2: 1.2 USULAN SOLUSI & 1.3 PEMANFAATAN AI (Google GenAI & RAG Vertex AI Search)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "1.2", "Usulan Solusi Berbasis Kecerdasan Artifisial")

add_paragraph(doc, [
    ("Sebagai jawaban atas tantangan manajemen data lingkungan hidup di Indonesia, platform ", False, False),
    ("Genesis", True, False),
    (" dirancang sebagai ekosistem aksi iklim lokal berbasis kecerdasan artifisial. Genesis mengintegrasikan aplikasi mobile Flutter (BLoC) "
     "untuk pelaporan warga, dasbor admin Next.js untuk koordinasi instansi, dan NestJS (GCP) untuk backend API. Platform ini mendefinisikan "
     "ulang sistem pelaporan pasif menjadi dinas aktif dengan menerapkan validasi spasial, reduksi data privasi, klasifikasi otomatis "
     "jenis kerusakan lingkungan, serta asisten regulasi cerdas berbasis RAG untuk meningkatkan kapasitas pengambilan keputusan Dinas Lingkungan Hidup.", False, False)
])

add_heading(doc, "1.3", "Pemanfaatan Kecerdasan Artifisial")

add_paragraph(doc, [
    ("Pemanfaatan AI pada Genesis diintegrasikan via ", False, False),
    ("Google GenAI SDK (Vertex AI)", True, False),
    (" untuk efisiensi operasional. AI mengklasifikasi jenis kerusakan lingkungan dan limbah secara otomatis menggunakan model ", False, False),
    ("Gemini 3.5 Flash", True, False),
    (" ke dalam format JSON terstruktur untuk sorting prioritas penanganan. Demi privasi warga, ", False, False),
    ("Google Vision API", True, False),
    (" mendeteksi wajah/plat nomor pada citra, lalu diburamkan secara in-memory via ", False, False),
    ("Sharp", True, False),
    (" sebelum disimpan di GCS. Asisten regulasi lingkungan diintegrasikan secara langsung menggunakan ", False, False),
    ("Vertex AI Search RAG", True, False),
    (" untuk mengakses dokumen hukum resmi, sedangkan fitur kueri suara didukung transkripsi audio otomatis melalui kemampuan ", False, False),
    ("native Speech-to-Text", False, True),
    (" bawaan dari model Gemini.", False, False)
])

add_paragraph(doc, [
    ("Guna mengoptimalkan latensi server dan efisiensi biaya pada Google Cloud lewat platform agen AI (", False, False),
    ("AI agent platform", False, True),
    ("), sistem menerapkan pola perutean model AI dinamis (", False, False),
    ("dynamic model routing", False, True),
    ("). Kueri chatbot sederhana diproses otomatis oleh model berbiaya rendah dengan latensi ultra-rendah yaitu ", False, False),
    ("Gemini 2.5 Flash Lite", True, False),
    (". Sebaliknya, kueri kompleks yang membutuhkan penalaran regulasi hukum panjang, dokumen pembanding RAG, atau klasifikasi foto kerusakan lingkungan diarahkan ke model dengan kapasitas penalaran tinggi yaitu ", False, False),
    ("Gemini 3.5 Flash", True, False),
    (". Pola perutean ini menjamin efisiensi biaya operasional dan kecepatan respons sistem.", False, False)
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PAGE 3: 1.4 PENDEKATAN TEKNIS DAN SUMBER DATA (Tabel 1 & 2)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "1.4", "Pendekatan Teknis dan Sumber Data")

add_paragraph(doc, [
    ("Genesis menggunakan arsitektur multiplatform: Flutter (BLoC) untuk mobile, Next.js untuk dasbor admin, dan NestJS (GCP) "
     "untuk backend API. Tim memanfaatkan ", False, False),
    ("Google Antigravity sebagai AI coding assistant", True, False),
     (" dalam optimasi kode dan dokumentasi, dengan perancangan arsitektur teruji secara mandiri. "
      "Sistem menerapkan otomasi pemrosesan data lingkungan terpadu dari hulu ke hilir yang dirangkum pada Tabel 1.", False, False),
])

# ── Tabel 1: Alur Kerja Sistem (Input, Proses, Output) ──
build_simple_table(doc,
    ["No.", "Tahap", "Aktivitas Teknis", "Implementasi Teknologi"],
    [
        ("1",
         "Input Data",
         "Warga memotret objek laporan via mobile client, koordinat GPS terlampir otomatis.",
         "Sensor kamera dan sensor GPS mobile."),
        
        ("2",
         "Sanitasi Data",
         "Deduplikasi spasial (radius 50m). Sensor wajah/plat nomor in-memory.",
         "Postgres/PostGIS, Google Vision API, Sharp."),
        
        ("3",
         "Analisis AI",
         "AI mengklasifikasi jenis limbah/kerusakan lingkungan, tingkat keparahan, dan rekomendasi mitigasi.",
         "Gemini 3.5 Flash menghasilkan metadata JSON terstruktur."),
        
        ("4",
         "Output Data",
         "Peta laporan tersensor di dasbor admin, poin gamifikasi warga, dan chatbot RAG suara.",
         "Vertex AI Search, Gemini, dan Server-Sent Events (SSE).")
    ],
    [1.0, 2.5, 5.5, 5.0] # Total: 14.0 cm
)
add_caption(doc, "Tabel 1. Alur input, proses, dan output pemrosesan data AI pada Genesis.")

# ── Paragraf Alur Teknis ──
add_paragraph(doc, [
    ("Alur pelaporan terintegrasi otomatis: deteksi duplikasi spasial di PostGIS (radius 50m, jeda 12h), "
     "sensor wajah/plat nomor via Google Vision API dan Sharp, klasifikasi limbah via Gemini 3.5 Flash "
     "(skor < 85% divalidasi manual oleh dinas), serta konsultasi hukum berbasis RAG (Vertex AI Search, native Gemini Speech-to-Text, SSE) bagi warga. "
     "Klasifikasi sumber data dirangkum pada Tabel 2.", False, False)
])

# ── Tabel 2: Sumber Data & Batasan Data dengan Kolom NO ──
build_simple_table(doc,
    ["No.", "Kategori Data", "Format Data", "Sumber Perolehan", "Batasan dan Keamanan Data"],
    [
        ("1",
         "Partisipasi Publik",
         "GPS, JPEG/PNG, M4A",
         "Sensor perangkat mobile warga",
         "Data uji dari pengembang dan tidak menyimpan informasi identitas pribadi (Personally Identifiable Information, PII)."),
        
        ("2",
         "Regulasi Daerah",
         "PDF / Markdown",
         "Portal data terbuka pemerintah",
         "Menggunakan dokumen resmi Perda, UU, dan PP lingkungan hidup publik."),
        
        ("3",
         "Data Historis",
         "Relasional SQL",
         "PostgreSQL (Supabase)",
         "Data samaran username warga dan poin gamifikasi tanpa data sensitif.")
    ],
    [1.0, 2.8, 2.2, 3.5, 4.5] # Total: 14.0 cm
)
add_caption(doc, "Tabel 2. Klasifikasi sumber data dan batasan keamanan data Genesis.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PAGE 4: 1.5 RESPONSIBLE AI (Tabel 3) & 1.6 DAMPAK DAN IMPLEMENTASI (Padat)
# ══════════════════════════════════════════════════════════════
add_heading(doc, "1.5", "Responsible AI")

add_paragraph(doc, [
    ("Genesis menerapkan prinsip etika AI sejak awal perancangan sistem menggunakan pendekatan Safety by Design. "
     "Distribusi implementasi nyata empat pilar Responsible AI dirangkum pada Tabel 3.", False, False)
])

# ── Tabel 3: Pilar Responsible AI dengan Kolom NO ──
build_simple_table(doc,
    ["No.", "Pilar Responsible AI", "Implementasi Nyata dan Mitigasi Risiko"],
    [
        ("1",
         "Privasi & Keamanan",
         "Pemburaman wajah dan plat nomor diproses langsung di RAM server sebelum disimpan. Token JWT membatasi akses data antara warga dan administrator."),
        
        ("2",
         "Transparansi & Rujukan",
         "Setiap hasil analisis AI disertai skor keyakinan. Chatbot regulasi menyertakan kutipan pasal resmi untuk menghindari kesalahan informasi hukum."),
        
        ("3",
         "Peran Manusia (HITL)",
         "AI sebagai asisten. Laporan dengan skor keyakinan di bawah 85% divalidasi oleh petugas Dinas Lingkungan Hidup secara manual sebelum disetujui."),
        
        ("4",
         "Keandalan Sistem",
         "Rate limiting membatasi frekuensi kueri chat. Input teks melalui chatbot divalidasi menggunakan sanitasi masukan dan deteksi pola prompt injection.")
    ],
    [1.0, 3.5, 9.5] # Total: 14.0 cm
)
add_caption(doc, "Tabel 3. Matriks implementasi pilar Responsible AI Genesis.")

# ── Paragraf Penutup 1.5 ──
add_paragraph(doc, [
    ("Genesis memitigasi risiko bias klasifikasi (verifikasi manual), misinformasi regulasi (pembatasan data Vertex AI Search), dan prompt injection (sanitasi input). "
     "AI berfungsi sebagai sistem pendukung keputusan (decision support system), sedangkan seluruh keputusan publik tetap berada pada kewenangan "
     "petugas Dinas Lingkungan Hidup sebagai pengambil keputusan akhir. Pendekatan ini memastikan AI berperan sebagai alat bantu analisis, bukan pengganti pertimbangan manusia.", False, False)
])

# ── 1.6 DAMPAK DAN IMPLEMENTASI ──
add_heading(doc, "1.6", "Dampak dan Implementasi")

add_paragraph(doc, [
    ("Implementasi platform Genesis di tingkat lokal diharapkan memberikan dampak nyata: "
     "(1) Dampak Lingkungan: Gamifikasi (XP dan koin) meningkatkan partisipasi aktif masyarakat dalam pemeliharaan lingkungan, mempercepat penanganan kerusakan ekologis, dan menurunkan emisi karbon lokal secara terukur. "
     "(2) Efisiensi Birokrasi: Algoritma deduplikasi spasial PostGIS dan klasifikasi Gemini 3.5 Flash mereduksi beban verifikasi administratif Dinas Lingkungan Hidup hingga 70%, mengoptimalkan pengerahan armada kebersihan. "
     "(3) Kemitraan & Keberlanjutan: Koin virtual warga dapat ditukarkan dengan insentif ekonomi nyata (voucher belanja, sembako) melalui kemitraan bank daur ulang dan ritel swasta lokal, membangun sistem pengambilan keputusan berbasis bukti.", False, False)
])

# ══════════════════════════════════════════════════════════════
# SIMPAN DENGAN GRACEFUL INCREMENTAL LOCK HANDLING
# ══════════════════════════════════════════════════════════════
saved = False
counter = 1
temp_out = OUTPUT
while not saved:
    try:
        doc.save(temp_out)
        print(f"OK: {temp_out}")
        saved = True
    except PermissionError:
        temp_out = OUTPUT.replace(".docx", f"_{counter}.docx")
        counter += 1
        if counter > 100: # prevent infinite loop
            print("ERROR: Gagal menyimpan karena seluruh nama file alternatif terkunci.")
            break
