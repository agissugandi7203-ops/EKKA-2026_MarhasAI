"""
Generator Word Document — WorldSkills / LKS Test Project Proposal (Premium Indonesian Version)
Kecerdasan Artifisial (Artificial Intelligence) LKS Nasional 2026
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = r"d:\PROJECT ARIEF\LKS Dikdasmen\docs\Test_Project_AI_Genesis.docx"
IMAGE_PATH = r"C:\Users\arief\.gemini\antigravity\brain\5d19354b-3bf1-42bf-bfc4-06f605652364\technical_architecture_diagram_1782888262343.png"

# Konstanta Format Dokumen
FONT_NAME = "Arial"
SIZE_TITLE = Pt(20)
SIZE_H1 = Pt(14)
SIZE_H2 = Pt(12)
SIZE_BODY = Pt(10.5)
SIZE_TABLE_TEXT = Pt(9.5)
LINE_SPACING_BODY = 1.15
LINE_SPACING_TABLE = 1.05

COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY = RGBColor(0x55, 0x55, 0x55)
COLOR_WS_BLUE = RGBColor(0x0A, 0x22, 0x40) # WorldSkills Dark Blue

HEX_BORDER = "CCCCCC"
HEX_HEADER_BG = "F2F2F2"

def style_cell(cell, top_padding=100, bottom_padding=100, left_padding=120, right_padding=120, 
               shading_color=None, border_color=HEX_BORDER, v_align="center"):
    """Mengatur padding internal sel, shading, border grid tipis, dan perataan vertikal."""
    tcPr = cell._tc.get_or_add_tcPr()
    
    # 1. Padding Sel
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top_padding}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom_padding}" w:type="dxa"/>'
        f'  <w:left w:w="{left_padding}" w:type="dxa"/>'
        f'  <w:right w:w="{right_padding}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)
    
    # 2. Shading Sel
    if shading_color:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
        tcPr.append(shd)
        
    # 3. Border Grid Sederhana (Semua sisi)
    for el in list(tcPr):
        if el.tag.endswith('tcBorders'):
            tcPr.remove(el)
    borders_xml = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders_xml)
    
    # 4. Perataan Vertikal Tengah
    for el in list(tcPr):
        if el.tag.endswith('vAlign'):
            tcPr.remove(el)
    v_align_xml = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{v_align}"/>')
    tcPr.append(v_align_xml)

def write_cell(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING_TABLE
    
    words = text.split(" ")
    for word in words:
        run = p.add_run(word + " ")
        run.font.name = FONT_NAME
        run.font.size = SIZE_TABLE_TEXT
        run.font.color.rgb = COLOR_BLACK
        run.bold = bold
        
        is_english = italic or any(term in word.lower() for term in [
            "multipart", "in-memory", "blur", "embedding", "streaming", "routing", 
            "check_duplicate_report", "pgvector", "hnsw", "whisper", "client", "record", 
            "geolocator", "database", "vision", "sharp", "flash", "openrouter", "throttler", 
            "jwt", "rbac", "rate", "limiting", "injection", "evasion", "typoglycemia", 
            "active", "learning", "feedback", "loop", "upvote", "design", "hitl", "gateway",
            "support", "similarity", "pending_human"
        ])
        run.italic = is_english

        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
            rPr.append(rFonts)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    pf.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = SIZE_H1
    run.font.color.rgb = COLOR_WS_BLUE
    run.bold = True
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    pf.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = SIZE_H2
    run.font.color.rgb = COLOR_BLACK
    run.bold = True
    return p

def add_paragraph(doc, text_segments, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = LINE_SPACING_BODY
    
    for txt, b, i in text_segments:
        run = p.add_run(txt)
        run.font.name = FONT_NAME
        run.font.size = SIZE_BODY
        run.font.color.rgb = COLOR_BLACK
        run.bold = b
        run.italic = i
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(9.0)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True

def add_centered_picture(doc, path, width):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=width)
    return p

def build_simple_table(doc, headers, data, col_widths):
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Header Row
    for col_idx, text in enumerate(headers):
        cell = tbl.rows[0].cells[col_idx]
        write_cell(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_cell(cell, shading_color=HEX_HEADER_BG)
    
    # Data Rows
    for row_idx, row_data in enumerate(data):
        row = tbl.rows[row_idx + 1]
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            bold = (col_idx == 0)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cell, val, bold=bold, align=align)
            style_cell(cell)
                
    # Atur Lebar Kolom (Total 15.5 cm agar pas dengan margin 3.0-2.5)
    for row in tbl.rows:
        for col_idx, width in enumerate(col_widths):
            row.cells[col_idx].width = Cm(width)
            
    return tbl

# ══════════════════════════════════════════════════════════════
# DOKUMEN SETUP (Margin WorldSkills/LKS: Left 3cm, Right 2.5cm)
# ══════════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.5)

# ── Konfigurasi Header dan Footer ──
header = sec.header
p_header = header.paragraphs[0]
p_header.text = ""
p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
h_run = p_header.add_run("LKS Nasional 2026 | Kecerdasan Artifisial (Artificial Intelligence)")
h_run.font.name = FONT_NAME
h_run.font.size = Pt(8.5)
h_run.font.color.rgb = COLOR_GRAY

footer = sec.footer
p_footer = footer.paragraphs[0]
p_footer.text = ""
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
f_run = p_footer.add_run("Tanggal: 05.07.2026  |  Versi: 1.0  |  LKS2026_TP_AI_ID  |  © WorldSkills International")
f_run.font.name = FONT_NAME
f_run.font.size = Pt(8.5)
f_run.font.color.rgb = COLOR_GRAY

# ── Halaman Judul Formal (Cover Page Style) ──
p_space = doc.add_paragraph()
p_space.paragraph_format.space_before = Pt(36)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title_format = p_title.paragraph_format
p_title_format.space_after = Pt(12)

run_lks = p_title.add_run("LOMBA KOMPETENSI SISWA (LKS) PENDIDIKAN MENENGAH\nTINGKAT NASIONAL XXXIV TAHUN 2026\n\n")
run_lks.font.name = FONT_NAME
run_lks.font.size = Pt(13)
run_lks.font.color.rgb = COLOR_GRAY
run_lks.bold = True

run_main_title = p_title.add_run("PROPOSAL PROYEK UJI (TEST PROJECT PROPOSAL)\nKecerdasan Artifisial (Artificial Intelligence)\n")
run_main_title.font.name = FONT_NAME
run_main_title.font.size = SIZE_TITLE
run_main_title.font.color.rgb = COLOR_WS_BLUE
run_main_title.bold = True

p_desc = doc.add_paragraph()
p_desc_format = p_desc.paragraph_format
p_desc_format.space_after = Pt(48)
run_desc = p_desc.add_run(
    "Nama Proyek: Genesis — Platform Pelaporan Ekologi Berbasis Crowdsourcing dengan Anti-Spam Spasial, "
    "Sensor Data Sensitif (PII Redaction), dan Chatbot RAG Hukum Berbasis Model Routing Hybrid\n"
)
run_desc.font.name = FONT_NAME
run_desc.font.size = Pt(11)
run_desc.italic = True
run_desc.font.color.rgb = COLOR_GRAY

# Page Break
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DAFTAR ISI
# ══════════════════════════════════════════════════════════════
add_heading_1(doc, "Daftar Isi")
add_paragraph(doc, [
    ("Proposal proyek uji ini terbagi atas beberapa bagian wajib sesuai ketentuan WorldSkills International:\n", False, False),
    ("1. Pendahuluan\n", True, False),
    ("   - Memuat latar belakang masalah ekologi, solusi konsep Genesis, dan bagan alur arsitektur teknis.\n", False, False),
    ("2. Deskripsi Proyek dan Tugas\n", True, False),
    ("   - Menjelaskan pembagian modul evaluasi (Modul A hingga Modul D) serta target hasil pembelajaran.\n", False, False),
    ("3. Instruksi kepada Peserta\n", True, False),
    ("   - Menyajikan spesifikasi teknis, hasil yang diharapkan, serta batasan untuk setiap modul kompetensi uji.\n", False, False),
    ("4. Lain-lain\n", True, False),
    ("   - Menjelaskan kerangka kerja Responsible AI, batasan data privasi, dan matriks penilaian LKS Nasional.", False, False),
], space_after=Pt(12))

# ══════════════════════════════════════════════════════════════
# PENDAHULUAN
# ══════════════════════════════════════════════════════════════
add_heading_1(doc, "Pendahuluan")
add_paragraph(doc, [
    ("Kompleksitas kerusakan akibat perubahan iklim di Indonesia memerlukan sistem pengambilan keputusan yang adaptif dan berbasis bukti. "
     "Platform pelaporan lingkungan konvensional saat ini sering kali mengalami kegagalan adopsi karena ketiadaan insentif partisipasi bagi warga, "
     "serta menyimpan data laporan sebagai entitas terisolasi yang berdiri sendiri. Ketiadaan integrasi data ini menghambat akumulasi wawasan ekologi lokal, "
     "sehingga memicu laporan ganda yang berulang dan membebani kinerja petugas administratif Dinas Lingkungan Hidup.\n\n"
     "Genesis dirancang sebagai platform aksi iklim lokal berbasis crowdsourcing untuk menjembatani celah tersebut. "
     "Sistem ini mengintegrasikan aplikasi mobile warga untuk pelaporan, dasbor admin web untuk kurasi data, "
     "serta server backend yang menerapkan ekstensi basis data spasial dan mesin AI pada Google Cloud Platform (GCP). Selama proses pengembangan, "
     "tim memanfaatkan Google Antigravity sebagai AI coding assistant untuk penyusunan dokumentasi teknis dan optimasi kode, "
     "sedangkan seluruh logika sistem, validasi pengujian, dan desain arsitektur diselesaikan secara mandiri oleh tim.\n\n"
     "Guna mengoptimalkan latensi server dan efisiensi biaya, sistem menerapkan pola perutean model AI dinamis (dynamic model routing). "
     "Kueri percakapan chatbot yang sederhana diproses secara otomatis oleh model berbiaya rendah dengan latensi ultra-rendah yaitu Gemini 2.5 Flash Lite. "
     "Sebaliknya, kueri kompleks yang membutuhkan penalaran regulasi hukum panjang, dokumen pembanding RAG, atau klasifikasi foto sampah "
     "menjadi metadata JSON diarahkan ke model dengan kapasitas penalaran tinggi yaitu Gemini 3.5 Flash. Pendekatan perutean model dinamis ini "
     "menghasilkan efisiensi sumber daya dan keamanan data, membentuk landasan yang kokoh untuk aksi iklim lokal.", False, False)
])

# Embed Architecture Diagram
if os.path.exists(IMAGE_PATH):
    add_centered_picture(doc, IMAGE_PATH, width=Cm(13.5))
    add_caption(doc, "Gambar 1. Diagram arsitektur teknis dan alur data AI pada platform Genesis.")

# Page Break
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DESKRIPSI PROYEK DAN TUGAS
# ══════════════════════════════════════════════════════════════
add_heading_1(doc, "Deskripsi Proyek dan Tugas")
add_paragraph(doc, [
    ("Proyek uji ini dirancang untuk menguji kompetensi peserta dalam integrasi Kecerdasan Artifisial, "
     "pengembangan full-stack, dan rekayasa data. Peserta wajib menyelesaikan empat modul kompetensi utama:", False, False),
])

add_heading_2(doc, "Modul A: Client Mobile & Inti Gamifikasi")
add_paragraph(doc, [
    ("Peserta wajib membangun aplikasi mobile warga menggunakan Flutter dengan arsitektur bersih dan manajemen status BLoC. "
     "Modul ini berfokus pada implementasi antarmuka pengguna (UI/UX), penangkapan koordinat GPS otomatis via geolocator saat memotret, "
     "serta visualisasi poin pengalaman (XP), level warga, lencana pencapaian (badges), dan papan peringkat (leaderboard) secara real-time.", False, False),
])

add_heading_2(doc, "Modul B: API Aman & Lapisan Akses Relasional")
add_paragraph(doc, [
    ("Peserta wajib membangun endpoint REST API menggunakan NestJS Fastify dan database relasional. "
     "Modul ini mencakup sistem autentikasi berbasis token JWT, otorisasi berbasis peran (RBAC) untuk menu warga (citizen) "
     "dan pengelola (admin), serta struktur tabel relasional PostgreSQL di Supabase.", False, False),
])

add_heading_2(doc, "Modul C: Pemrosesan Spasial & Sanitasi Computer Vision")
add_paragraph(doc, [
    ("Peserta wajib mengonfigurasi algoritma spasial dan otomatisasi sanitasi gambar. "
     "Modul ini membutuhkan fungsi database PostGIS check_duplicate_report() untuk mendeteksi laporan aktif dalam radius 50 meter "
     "dan jendela waktu 12 jam agar laporan baru otomatis digabungkan sebagai upvote. Selain itu, peserta wajib menggunakan model computer vision "
     "untuk mendeteksi wajah dan plat nomor kendaraan secara otomatis, lalu memburamkannya secara in-memory menggunakan Sharp sebelum diunggah ke Google Cloud Storage (GCS).", False, False),
])

add_heading_2(doc, "Modul D: Retrieval-Augmented Generation (RAG) & Perutean Model AI Dinamis")
add_paragraph(doc, [
    ("Peserta wajib mengimplementasikan sistem pencarian semantik dan perutean model AI dinamis. "
     "Modul ini mencakup pencarian kesamaan kosinus pada database vektor pgvector (indeks HNSW), transkripsi audio chatbot menggunakan Whisper-1, "
     "pembuatan API Router untuk membagi kueri pendek ke Gemini 2.5 Flash Lite dan kueri panjang/RAG ke Gemini 3.5 Flash, serta pengiriman data streaming via SSE.", False, False),
])

add_heading_2(doc, "Target Capaian Pembelajaran (Expected Learning Outcomes)")
add_paragraph(doc, [
    ("Dengan menyelesaikan proyek uji ini, peserta membuktikan penguasaan dalam domain teknis berikut:\n"
     "1. Integrasi AI: Penerapan praktis model LLM, embedding, dan pengolahan computer vision.\n"
     "2. Computer Vision: Proteksi privasi data warga via sensor otomatis wajah dan plat nomor kendaraan.\n"
     "3. Pemrosesan Geospasial: Deduplikasi spasial berbasis radius dan rentang waktu.\n"
     "4. Information Retrieval: Pencarian semantik berbasis vektor dan RAG regulasi hukum.\n"
     "5. Full-Stack Engineering: Desain API backend yang aman, database relasional, dan aplikasi mobile multiplatform.\n"
     "6. Responsible AI: Penerapan pengolahan data aman, pembatasan prompt injection, dan batas verifikasi manual manusia (HITL).", False, False),
])

# ══════════════════════════════════════════════════════════════
# INSTRUKSI KEPADA PESERTA
# ══════════════════════════════════════════════════════════════
add_heading_1(doc, "Instruksi kepada Peserta")

# ── Modul A ──
add_heading_2(doc, "Modul A: Client Mobile & Inti Gamifikasi")
add_paragraph(doc, [
    ("Tujuan: Mengimplementasikan aplikasi mobile warga untuk pelaporan dan umpan balik gamifikasi.\n", True, False),
    ("Persyaratan:\n", True, False),
    ("- Peserta wajib membuat halaman pelaporan sampah yang mengakses kamera dan geolocator perangkat.\n"
     "- Peserta wajib menggunakan pola BLoC untuk mengelola status perolehan poin XP, level warga, dan lencana.\n", False, False),
    ("Hasil yang Diharapkan: ", True, False),
    ("UI aplikasi mobile fungsional yang menampilkan metadata lokasi laporan dan data poin gamifikasi.\n", False, False),
    ("Batasan: ", True, False),
    ("Aplikasi harus membaca koordinat lokasi secara otomatis tanpa memerlukan input teks manual dari warga.", False, False)
])

# ── Modul B ──
add_heading_2(doc, "Modul B: API Aman & Lapisan Akses Relasional")
add_paragraph(doc, [
    ("Tujuan: Membangun endpoint REST API yang aman dan skema database relasional.\n", True, False),
    ("Persyaratan:\n", True, False),
    ("- Peserta wajib mengimplementasikan API controller yang dilindungi autentikasi JWT token.\n"
     "- Peserta wajib menerapkan sistem otorisasi peran (RBAC) untuk mengamankan menu admin dari akses warga.\n", False, False),
    ("Hasil yang Diharapkan: ", True, False),
    ("Endpoint API yang terdokumentasi di Swagger dan terhubung ke database relasional.\n", False, False),
    ("Batasan: ", True, False),
    ("Seluruh akses pelaporan harus terautentikasi. Permintaan anonim wajib ditolak sistem.", False, False)
])

# ── Modul C ──
add_heading_2(doc, "Modul C: Pemrosesan Spasial & Sanitasi Computer Vision")
add_paragraph(doc, [
    ("Tujuan: Mengimplementasikan pencegahan laporan spam spasial dan sensor data sensitif otomatis.\n", True, False),
    ("Persyaratan:\n", True, False),
    ("- Peserta wajib membuat fungsi spasial SQL menggunakan ST_DWithin untuk mencari duplikasi dalam radius 50m.\n"
     "- Peserta wajib mengintegrasikan model computer vision untuk mendeteksi koordinat wajah/plat nomor, dan menerapkan filter Gaussian blur via Sharp.\n", False, False),
    ("Hasil yang Diharapkan: ", True, False),
    ("Foto yang tersimpan di cloud storage terbukti tersensor, dan laporan dalam radius 50m tergabung otomatis.\n", False, False),
    ("Batasan: ", True, False),
    ("Proses pemotongan/sensor gambar wajib dikerjakan in-memory di RAM server. Menyimpan gambar mentah tanpa blur adalah pelanggaran keamanan.", False, False)
])

# ── Modul D ──
add_heading_2(doc, "Modul D: Retrieval-Augmented Generation (RAG) & Perutean Model AI Dinamis")
add_paragraph(doc, [
    ("Tujuan: Membangun asisten chatbot regulasi hukum berbasis RAG dengan perutean model AI dinamis.\n", True, False),
    ("Persyaratan:\n", True, False),
    ("- Peserta wajib menerapkan pencarian semantik menggunakan kesamaan kosinus pada database vektor.\n"
     "- Peserta wajib merancang mekanisme API Router: mengarahkan kueri percakapan pendek (< 100 karakter) ke model Gemini 2.5 Flash Lite, dan kueri analisis dokumen/RAG panjang ke model Gemini 3.5 Flash.\n", False, False),
    ("Hasil yang Diharapkan: ", True, False),
    ("Chatbot regulasi lingkungan hidup yang merespons secara streaming menggunakan Server-Sent Events (SSE).\n", False, False),
    ("Batasan: ", True, False),
    ("Respons chatbot wajib bersumber secara mutlak pada dokumen regulasi resmi yang tersimpan di database vektor.", False, False)
])

# Page Break
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# LAIN-LAIN
# ══════════════════════════════════════════════════════════════
add_heading_1(doc, "Lain-lain")

add_heading_2(doc, "1. Penerapan Responsible AI dan Batasan Keamanan")
add_paragraph(doc, [
    ("Peserta diwajibkan menerapkan prinsip etika kecerdasan artifisial secara nyata pada seluruh deliverables:\n"
     "- Privasi (PII Redaction): Sensor wajah dan plat nomor diproses secara in-memory untuk mematuhi regulasi perlindungan data pribadi.\n"
     "- Keandalan (Prompt Injection Guardrails): Backend mendeteksi dan menolak pola prompt injection sebelum dikirim ke mesin LLM.\n"
     "- Pengawasan Manusia (Human-in-the-Loop): AI diposisikan sebagai sistem pendukung keputusan (decision support system). Laporan dengan skor keyakinan AI di bawah 85% otomatis disetel ke status \"pending_human\" untuk divalidasi manual oleh petugas Dinas Lingkungan Hidup melalui dasbor Next.js.", False, False)
])

add_heading_2(doc, "2. Batasan Data Uji")
add_paragraph(doc, [
    ("Seluruh pengembangan proyek uji hanya menggunakan data yang aman dan bebas data PII:\n"
     "- Regulasi Publik: Peraturan daerah, Peraturan Presiden, dan Undang-Undang lingkungan hidup resmi yang tersedia publik.\n"
     "- Dataset Terbuka: Kumpulan citra sampah publik untuk pelatihan model klasifikasi.\n"
     "- Data Sintetis: Simulasi data profil warga samaran, koordinat GPS pengembang, dan pencatatan poin gamifikasi.", False, False)
])

add_heading_2(doc, "3. Rubrik Penilaian Proyek Uji")
add_paragraph(doc, [
    ("Matriks rubrik penilaian juri nasional LKS 2026 untuk menguji performa kompetitor didistribusikan secara transparan pada Tabel 2.", False, False)
])

# ── Tabel Rubrik Penilaian Proyek ──
build_simple_table(doc,
    ["No.", "Kriteria Evaluasi Juri", "Bobot", "Metode Pengujian & Aspek Penilaian"],
    [
        ("1", "Pemahaman masalah & relevansi case", "20%", "Kesesuaian solusi dengan kebutuhan aksi iklim lokal, reduksi spam spasial, dan retensi pengguna."),
        ("2", "Kreativitas & inovasi solusi", "20%", "Keunikan integrasi gamifikasi warga, deduplikasi spasial PostGIS, dan perutean model AI dinamis."),
        ("3", "Pemanfaatan AI yang efektif & berdampak", "20%", "Akurasi model visi AI dalam klasifikasi sampah dan performa chatbot RAG regulasi."),
        ("4", "Penerapan Responsible AI (HITL & Privasi)", "15%", "Keamanan pemrosesan sensor gambar in-memory dan keandalan batas status pending_human."),
        ("5", "Fungsionalitas aplikasi/prototipe", "15%", "Kestabilan instalasi Flutter client, NestJS REST API, dan visualisasi admin dashboard."),
        ("6", "Kejelasan presentasi & dokumentasi", "10%", "Kerapian penulisan kode (clean code), dokumentasi teknis proposal, dan video demo.")
    ],
    [1.0, 5.5, 1.5, 7.5] # Total: 15.5 cm
)
add_caption(doc, "Tabel 2. Matriks kriteria penilaian juri ekshibisi AI LKS Nasional 2026.")

# ══════════════════════════════════════════════════════════════
# SIMPAN DOKUMEN
# ══════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"OK: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT)/1024:.1f} KB")
